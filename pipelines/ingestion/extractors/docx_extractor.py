from docx import Document
from docx.oxml.ns import qn
from typing import List
from pipelines.ingestion.base import ExtractedPage


def extract_docx(file_bytes: bytes) -> List[ExtractedPage]:
    pages: List[ExtractedPage] = []
    doc = Document(file_bytes)
    current_section = 1
    current_heading_level = None

    # Extract paragraphs
    for para in doc.paragraphs:
        metadata = {"section": current_section}
        
        # Check if it's a heading
        if para.style.name.startswith("Heading"):
            try:
                level = int(para.style.name.split()[-1])
                current_heading_level = level
                metadata["heading_level"] = level
                metadata["heading_text"] = para.text
            except ValueError:
                pass
        
        if para.text.strip():
            pages.append(ExtractedPage(
                page_number=current_section,  # DOCX doesn't have real page numbers
                content=para.text,
                content_type="text",
                metadata=metadata
            ))

    # Extract tables
    for table_idx, table in enumerate(doc.tables):
        # Convert table to markdown
        md = []
        for row_idx, row in enumerate(table.rows):
            cleaned_row = [cell.text.strip() for cell in row.cells]
            md_row = "| " + " | ".join(cleaned_row) + " |"
            md.append(md_row)
            if row_idx == 0:
                md.append("| " + " | ".join(["---"] * len(cleaned_row)) + " |")
        table_md = "\n".join(md)
        pages.append(ExtractedPage(
            page_number=0,
            content=table_md,
            content_type="table",
            metadata={"table_index": table_idx}
        ))
    
    return pages
